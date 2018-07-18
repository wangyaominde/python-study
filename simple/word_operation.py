from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

Doc = Document()

'''p = Doc.add_paragraph("昆明沃创科技有限公司")
p1 = Doc.add_paragraph("培训通知")
p2 = Doc.add_paragraph("沃创科技〔2017〕 号")

p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
'''
p = Doc.add_paragraph()
run = p.add_run("昆明沃创科技有限公司")
run.font.color.rgb = RGBColor(255,0,0)


Doc.save('test.docx')
