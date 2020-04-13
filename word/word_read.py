#读取docx中的文本代码示例
import docx
#获取文档对象
file=docx.Document("./1912SX02HK01051.docx")
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
'''
#输出每一段的内容
for para in file.paragraphs:
    print(para.text)
'''
#输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    if(i==4):
        print(file.paragraphs[i].text[6:])
